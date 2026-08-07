"""Builds an editable .docx document from per-page OCR results."""

from __future__ import annotations

from typing import List, Optional

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from .ocr_engine import OCRPageResult

# Vertical gap (in pixels, scaled by DPI) larger than this between two lines
# is treated as a paragraph break rather than a line-wrap within one paragraph.
PARAGRAPH_GAP_RATIO = 1.6


class DocxBuilder:
    def __init__(self, font_name: str = "Calibri", font_size: int = 11):
        self.document = Document()
        style = self.document.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)

    def add_page(
        self,
        page_result: Optional[OCRPageResult],
        page_index: int,
        native_text: Optional[str] = None,
        detected_language: Optional[str] = None,
    ) -> None:
        """Append one page's worth of content to the document."""
        if page_index > 0:
            self._add_page_break()

        if detected_language:
            note = self.document.add_paragraph()
            run = note.add_run(f"[Page {page_index + 1} - detected language: {detected_language}]")
            run.italic = True
            run.font.size = Pt(8)

        if native_text and native_text.strip():
            self._add_native_text(native_text)
            return

        self._add_ocr_lines(page_result)

    def _add_native_text(self, text: str) -> None:
        for block in text.split("\n\n"):
            block = block.strip()
            if block:
                self.document.add_paragraph(block)

    def _add_ocr_lines(self, page_result: OCRPageResult) -> None:
        if not page_result.lines:
            self.document.add_paragraph("[No text recognized on this page]")
            return

        avg_line_height = self._average_line_height(page_result)
        paragraph = self.document.add_paragraph()
        prev_bottom = None

        for line in page_result.lines:
            top = line.bbox[1]
            if prev_bottom is not None and avg_line_height:
                gap = top - prev_bottom
                if gap > avg_line_height * PARAGRAPH_GAP_RATIO:
                    paragraph = self.document.add_paragraph()
                elif paragraph.runs:
                    paragraph.add_run(" ")

            paragraph.add_run(line.text)
            prev_bottom = line.bbox[1] + line.bbox[3]

    @staticmethod
    def _average_line_height(page_result: OCRPageResult) -> float:
        heights = [line.bbox[3] for line in page_result.lines if line.bbox[3] > 0]
        return sum(heights) / len(heights) if heights else 0.0

    def _add_page_break(self) -> None:
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def save(self, output_path: str) -> None:
        self.document.save(output_path)
