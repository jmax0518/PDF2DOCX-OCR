"""Assemble an editable DOCX from PageLayout (text + images + tables)."""

from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from .layout_model import ImageBlock, PageLayout, TableBlock, TextBlock


class LayoutDocxBuilder:
    def __init__(self, font_name: str = "Calibri", default_font_size: int = 11, dpi: int = 300):
        self.document = Document()
        self.font_name = font_name
        self.default_font_size = default_font_size
        self.dpi = dpi
        style = self.document.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(default_font_size)
        self._first_page = True
        self._prev_bottom: Optional[float] = None
        self._page_height_px: float = 0.0

    def add_page(self, layout: PageLayout) -> None:
        if not self._first_page:
            self._add_page_break()
        self._first_page = False

        self._apply_page_size(layout)
        self._prev_bottom = None
        self._page_height_px = layout.height

        elements = layout.sorted_elements()
        if not elements:
            self.document.add_paragraph("")
            return

        for element in elements:
            self._apply_vertical_gap(element.bbox.y0, layout.height)
            if isinstance(element, TextBlock):
                self._add_text(element)
            elif isinstance(element, ImageBlock):
                self._add_image(element, layout.width)
            elif isinstance(element, TableBlock):
                self._add_table(element)
            self._prev_bottom = element.bbox.y1

    def save(self, output_path: str) -> None:
        self.document.save(output_path)

    def _apply_page_size(self, layout: PageLayout) -> None:
        section = self.document.sections[-1]
        # Convert pixel dimensions at pipeline DPI to Word inches.
        width_in = layout.width / float(self.dpi) if self.dpi else 8.5
        height_in = layout.height / float(self.dpi) if self.dpi else 11.0
        # Clamp to sane printable sizes
        width_in = max(4.0, min(width_in, 22.0))
        height_in = max(4.0, min(height_in, 22.0))
        section.page_width = Inches(width_in)
        section.page_height = Inches(height_in)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    def _apply_vertical_gap(self, next_top: float, page_height: float) -> None:
        if self._prev_bottom is None or page_height <= 0:
            return
        gap_px = next_top - self._prev_bottom
        if gap_px <= 0:
            return
        # Map pixel gap to paragraph spacing (points). ~dpi px = 72 pt.
        gap_pt = gap_px * 72.0 / float(self.dpi)
        # Cap spacer so large empty regions don't explode document length.
        gap_pt = max(0.0, min(gap_pt, 36.0))
        if gap_pt < 4:
            return
        spacer = self.document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(gap_pt)
        spacer.paragraph_format.line_spacing = 1.0
        # Keep an empty run so the paragraph is retained.
        spacer.add_run("")

    def _add_text(self, block: TextBlock) -> None:
        # Split on blank lines into paragraphs; keep single newlines as soft breaks.
        chunks = [c.strip() for c in block.text.split("\n\n") if c.strip()]
        if not chunks:
            return
        size = self.default_font_size
        if block.font_size_hint:
            size = int(round(max(8.0, min(block.font_size_hint, 36.0))))

        for chunk in chunks:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(chunk.replace("\n", " "))
            run.font.name = self.font_name
            run.font.size = Pt(size)
            # Ensure East Asian font fallback uses same name on Windows Word.
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), self.font_name)
            rfonts.set(qn("w:hAnsi"), self.font_name)

    def _add_image(self, block: ImageBlock, page_width_px: float) -> None:
        image = block.image
        if image.mode not in ("RGB", "RGBA", "L"):
            image = image.convert("RGB")

        # Target width: fraction of content width based on bbox vs page.
        section = self.document.sections[-1]
        content_width = section.page_width - section.left_margin - section.right_margin
        if page_width_px > 0:
            fraction = max(0.15, min(1.0, block.bbox.width / page_width_px))
        else:
            fraction = 0.8
        width = int(content_width) * fraction
        # python-docx accepts Inches; convert from EMUs (content_width is EMU).
        width_inches = width / 914400.0

        buffer = io.BytesIO()
        save_img = image.convert("RGB")
        save_img.save(buffer, format="PNG")
        buffer.seek(0)

        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run()
        run.add_picture(buffer, width=Inches(max(0.5, min(width_inches, 7.5))))

    def _add_table(self, block: TableBlock) -> None:
        rows = block.rows
        if not rows:
            return
        cols = max(len(r) for r in rows)
        if cols == 0:
            return
        table = self.document.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx in range(cols):
                value = row[c_idx] if c_idx < len(row) else ""
                table.cell(r_idx, c_idx).text = value
        # Spacer paragraph after table
        self.document.add_paragraph("")

    def _add_page_break(self) -> None:
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
