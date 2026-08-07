"""Exact-layout DOCX builder: place text/images at PDF page coordinates."""

from __future__ import annotations

import io
import re
from typing import Any, List, Optional, Tuple
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.shared import Inches, Pt
from PIL import Image

# Layout labels that should be rendered as pictures for visual fidelity.
_IMAGE_LABELS = frozenset(
    {"image", "chart", "seal", "figure", "table", "header_image", "footer_image"}
)

# Soft OCR cleanup: space after sentence-ending punctuation before a capital.
_SPACE_AFTER_PUNCT = re.compile(r"([.!?])([A-Z])")


def _px_to_inches(px: float, dpi: int) -> float:
    return float(px) / float(dpi)


def _inches_to_emu(inches: float) -> int:
    return int(round(inches * 914400))


def _inches_to_pt(inches: float) -> float:
    return inches * 72.0


def _clean_ocr_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _SPACE_AFTER_PUNCT.sub(r"\1 \2", text)
    return text.strip()


class ExactDocxBuilder:
    """Rebuild pages with absolute textboxes + floating images (visual fidelity)."""

    def __init__(self, dpi: int = 150, font_name: str = "Times New Roman"):
        self.dpi = dpi
        self.font_name = font_name
        self.document = Document()
        self._first_page = True
        self._shape_seq = 0
        self._clear_body()

    def _clear_body(self) -> None:
        body = self.document.element.body
        for child in list(body):
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)

    def add_page(self, page_result: Any, page_image: Image.Image) -> None:
        page_w_px = page_image.width
        page_h_px = page_image.height
        page_w_in = _px_to_inches(page_w_px, self.dpi)
        page_h_in = _px_to_inches(page_h_px, self.dpi)

        if self._first_page:
            self._apply_page_size(self.document.sections[0], page_w_in, page_h_in)
            self._first_page = False
        else:
            section = self.document.add_section()
            self._apply_page_size(section, page_w_in, page_h_in)

        blocks = list(page_result.get("parsing_res_list") or [])
        # Larger regions first so titles/captions stay above when overlapping.
        blocks.sort(key=lambda b: getattr(b, "area", 0) or 0, reverse=True)

        host = self.document.add_paragraph()
        host.paragraph_format.space_before = Pt(0)
        host.paragraph_format.space_after = Pt(0)

        for block in blocks:
            label = (getattr(block, "label", None) or "").lower()
            bbox = getattr(block, "bbox", None)
            if not bbox or len(bbox) < 4:
                continue
            x0, y0, x1, y1 = [float(v) for v in bbox[:4]]
            if x1 <= x0 or y1 <= y0:
                continue

            left_in = max(0.0, _px_to_inches(x0, self.dpi))
            top_in = max(0.0, _px_to_inches(y0, self.dpi))
            width_in = max(0.05, _px_to_inches(x1 - x0, self.dpi))
            height_in = max(0.05, _px_to_inches(y1 - y0, self.dpi))
            width_in = min(width_in, max(0.05, page_w_in - left_in))
            height_in = min(height_in, max(0.05, page_h_in - top_in))

            pil_img = self._block_image(block, page_image, (x0, y0, x1, y1), label)
            if pil_img is not None:
                self._add_floating_image(host, pil_img, left_in, top_in, width_in, height_in)
                continue

            text = getattr(block, "content", None) or ""
            if not str(text).strip():
                continue
            text = _clean_ocr_text(str(text))
            font_pt = self._estimate_font_pt(block, height_in)
            bold = label in {"doc_title", "paragraph_title", "figure_title"}
            self._add_textbox(
                host,
                text,
                left_in,
                top_in,
                width_in,
                height_in,
                font_pt=font_pt,
                bold=bold,
            )

    def save(self, output_path: str) -> None:
        self.document.save(output_path)

    @staticmethod
    def _apply_page_size(section, width_in: float, height_in: float) -> None:
        width_in = max(4.0, min(width_in, 22.0))
        height_in = max(4.0, min(height_in, 22.0))
        if width_in > height_in * 1.05:
            section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(width_in)
        section.page_height = Inches(height_in)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)

    def _estimate_font_pt(self, block: Any, height_in: float) -> float:
        line_h_px = getattr(block, "text_line_height", None)
        n_lines = getattr(block, "num_of_lines", None) or 1
        if line_h_px and line_h_px > 0:
            pt = float(line_h_px) * 72.0 / float(self.dpi)
        else:
            pt = (height_in * 72.0) / max(float(n_lines), 1.0)
        pt = pt * 0.82
        return max(6.0, min(pt, 28.0))

    def _block_image(
        self,
        block: Any,
        page_image: Image.Image,
        bbox: Tuple[float, float, float, float],
        label: str,
    ) -> Optional[Image.Image]:
        image_info = getattr(block, "image", None)
        if isinstance(image_info, dict) and image_info.get("img") is not None:
            img = image_info["img"]
            if isinstance(img, Image.Image):
                return img.convert("RGB")

        if label in _IMAGE_LABELS:
            return self._crop_page(page_image, bbox)
        return None

    @staticmethod
    def _crop_page(
        page_image: Image.Image, bbox: Tuple[float, float, float, float]
    ) -> Optional[Image.Image]:
        x0, y0, x1, y1 = bbox
        w, h = page_image.size
        ix0 = max(0, int(x0))
        iy0 = max(0, int(y0))
        ix1 = min(w, int(x1))
        iy1 = min(h, int(y1))
        if ix1 - ix0 < 4 or iy1 - iy0 < 4:
            return None
        return page_image.crop((ix0, iy0, ix1, iy1)).convert("RGB")

    def _next_shape_id(self) -> str:
        self._shape_seq += 1
        return f"_exact{self._shape_seq}"

    def _add_textbox(
        self,
        paragraph,
        text: str,
        left_in: float,
        top_in: float,
        width_in: float,
        height_in: float,
        font_pt: float,
        bold: bool,
    ) -> None:
        left_pt = _inches_to_pt(left_in)
        top_pt = _inches_to_pt(top_in)
        width_pt = _inches_to_pt(width_in)
        height_pt = _inches_to_pt(height_in)
        sz = int(round(font_pt * 2))
        bold_xml = "<w:b/>" if bold else ""
        font = escape(self.font_name)

        para_chunks: List[str] = []
        for chunk in text.split("\n"):
            t = escape(chunk)
            para_chunks.append(
                "<w:p>"
                '<w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
                '<w:ind w:left="0" w:right="0"/></w:pPr>'
                "<w:r>"
                f"<w:rPr>{bold_xml}<w:sz w:val=\"{sz}\"/><w:szCs w:val=\"{sz}\"/>"
                f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
                "</w:rPr>"
                f'<w:t xml:space="preserve">{t}</w:t>'
                "</w:r></w:p>"
            )
        content_xml = "".join(para_chunks) or "<w:p/>"
        shape_id = self._next_shape_id()

        pict_xml = f"""
        <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:v="urn:schemas-microsoft-com:vml"
                xmlns:o="urn:schemas-microsoft-com:office:office"
                xmlns:w10="urn:schemas-microsoft-com:office:word">
          <v:shape id="{shape_id}" type="#_x0000_t202"
            style="position:absolute;margin-left:{left_pt:.2f}pt;margin-top:{top_pt:.2f}pt;width:{width_pt:.2f}pt;height:{height_pt:.2f}pt;z-index:251659264;mso-wrap-style:square;mso-width-percent:0;mso-height-percent:0;mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page;"
            filled="f" stroked="f" o:allowincell="f">
            <v:textbox style="mso-fit-shape-to-text:f" inset="0,0,0,0">
              <w:txbxContent>
                {content_xml}
              </w:txbxContent>
            </v:textbox>
          </v:shape>
        </w:pict>
        """
        run = paragraph.add_run()
        run._r.append(parse_xml(pict_xml))

    def _add_floating_image(
        self,
        paragraph,
        image: Image.Image,
        left_in: float,
        top_in: float,
        width_in: float,
        height_in: float,
    ) -> None:
        buf = io.BytesIO()
        image.save(buf, format="PNG")
        buf.seek(0)

        run = paragraph.add_run()
        inline_shape = run.add_picture(buf, width=Inches(width_in), height=Inches(height_in))
        inline = inline_shape._inline
        drawing = inline.getparent()

        left_emu = _inches_to_emu(left_in)
        top_emu = _inches_to_emu(top_in)
        cx = int(inline.extent.cx)
        cy = int(inline.extent.cy)

        docPr_xml = inline.docPr.xml
        graphic = None
        for child in inline.iter():
            if child.tag.endswith("}graphic"):
                graphic = child
                break
        if graphic is None:
            raise RuntimeError("Could not locate DrawingML graphic for floating image.")
        graphic_xml = graphic.xml

        anchor_xml = f"""
        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                   xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                   xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                   distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="251659264" behindDoc="0" locked="0"
                   layoutInCell="1" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page"><wp:posOffset>{left_emu}</wp:posOffset></wp:positionH>
          <wp:positionV relativeFrom="page"><wp:posOffset>{top_emu}</wp:posOffset></wp:positionV>
          <wp:extent cx="{cx}" cy="{cy}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          {docPr_xml}
          <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks noChangeAspect="1"/>
          </wp:cNvGraphicFramePr>
          {graphic_xml}
        </wp:anchor>
        """
        anchor = parse_xml(anchor_xml)
        drawing.replace(inline, anchor)
